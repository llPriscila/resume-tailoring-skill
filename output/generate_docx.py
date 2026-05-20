from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

DARK = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x55, 0x55, 0x55)
LINK = RGBColor(0x05, 0x63, 0xC1)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)

    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    run_elem.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run_elem.append(t)
    hl.append(run_elem)
    paragraph._p.append(hl)


def set_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'cccccc')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_runs(paragraph, text, size, bold=False, italic=False, color=None):
    """Parse markdown inline (links, bold, italic) and add runs to paragraph."""
    text = text.replace('—', '–')
    pattern = r'\[([^\]]+)\]\(([^\)]+)\)|\*\*([^*]+)\*\*|\*([^*]+)\*'
    last = 0
    for m in re.finditer(pattern, text):
        if m.start() > last:
            run = paragraph.add_run(text[last:m.start()])
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = color
        link_text, link_url, bold_text, italic_text = m.groups()
        if link_text:
            add_hyperlink(paragraph, link_text, link_url)
        elif bold_text:
            run = paragraph.add_run(bold_text)
            run.font.size = Pt(size)
            run.font.bold = True
            if color:
                run.font.color.rgb = color
        elif italic_text:
            run = paragraph.add_run(italic_text)
            run.font.size = Pt(size)
            run.font.italic = True
            if color:
                run.font.color.rgb = color
        last = m.end()
    if last < len(text):
        run = paragraph.add_run(text[last:])
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color


def sp(paragraph, before=0, after=0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def build_docx(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()

    # Page setup: A4, tight margins
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)

    # Default style
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9.5)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # Skip horizontal rules
        if line.strip() in ('---', '___', '***'):
            i += 1
            continue

        # H1 — name
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_paragraph()
            sp(p, before=0, after=2)
            run = p.add_run(line[2:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = DARK
            i += 1
            continue

        # H2 — section header with bottom border
        if line.startswith('## '):
            p = doc.add_paragraph()
            sp(p, before=8, after=3)
            run = p.add_run(line[3:].strip().upper())
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = DARK
            set_bottom_border(p)
            i += 1
            continue

        # H3 — role title
        if line.startswith('### '):
            p = doc.add_paragraph()
            sp(p, before=6, after=0)
            add_runs(p, line[4:].strip(), size=10, bold=True, color=DARK)
            i += 1
            continue

        # Bold-only line — role meta (Company — Location | Dates)
        if (line.startswith('**') and line.endswith('**')
                and not re.match(r'^\*\*([^*]+)\*\*[:\s]+\S', line)):
            p = doc.add_paragraph()
            sp(p, before=0, after=3)
            add_runs(p, line[2:-2].strip(), size=9, color=GRAY)
            i += 1
            continue

        # Bullet point
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            sp(p, before=0, after=1)
            p.paragraph_format.left_indent = Cm(0.4)
            add_runs(p, line[2:].strip(), size=9.5, color=DARK)
            i += 1
            continue

        # Skills / label line: **Label:** content
        m = re.match(r'^\*\*([^*]+)\*\*[:\s]+(.*)', line)
        if m:
            label = m.group(1).strip().rstrip(':')
            rest = m.group(2).strip()
            p = doc.add_paragraph()
            sp(p, before=0, after=4)
            run = p.add_run(f'{label}: ')
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = DARK
            if label.lower().startswith('currently'):
                run.font.italic = True
                run.font.color.rgb = GRAY
                r2 = p.add_run(rest)
                r2.font.size = Pt(9)
                r2.font.italic = True
                r2.font.color.rgb = GRAY
            else:
                add_runs(p, rest, size=9.5, color=DARK)
            i += 1
            continue

        # Non-empty body / contact / summary
        if line.strip():
            p = doc.add_paragraph()
            sp(p, before=0, after=3)
            add_runs(p, line.strip(), size=9.5, color=DARK)
            i += 1
            continue

        # Empty line
        p = doc.add_paragraph()
        sp(p, before=0, after=2)
        i += 1

    doc.save(docx_path)
    print(f'Generated: {docx_path}')


def md_to_docx(md_filename, base_dir=None):
    base_dir = base_dir or os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base_dir, md_filename)
    docx_path = os.path.join(base_dir, os.path.splitext(md_filename)[0] + '.docx')
    build_docx(md_path, docx_path)
    return docx_path


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    for f in [
        'Ioana_Criclevit_Wise_Staff_Data_Analyst_Resume.md',
        'Ioana_Criclevit_Ravio_Insights_Analyst_Resume.md',
        'Ioana_Criclevit_FCA_Financial_Crime_Senior_Data_Analyst_Resume.md',
        'Ioana_Criclevit_Zopa_Data_Analyst_Fraud_Financial_Crime_Resume.md',
        'Ioana_Criclevit_Zopa_Senior_Analyst_Lending_Operations_Resume.md',
        'Ioana_Criclevit_Monzo_Senior_Data_Analyst_Financial_Health_Resume.md',
    ]:
        md_to_docx(f, base)
